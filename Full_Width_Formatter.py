r"""
Indentor v1 (Windows-only)

Batch-insert two full-width spaces (U+3000×2) at the start of each line (.txt) or each paragraph (.docx).

- Formats: .txt, .docx (NO .doc)
- Naming: if conflict → _indent; if conflict persists → _copy(2), _copy(3) ...
- Encoding (.txt): detect & preserve; preserves per-line line endings
- .docx: inserts after bullets/numbering; skips empty paragraphs; skips Heading styles
- UI: minimal, pretty; shows up to 3 selected files; shows processing & last 3 finished; has "Open output folder"
- Note line: "Please convert .doc to .docx before using this tool!"

Dependencies (install in Windows):
    pip install PySide6 python-docx charset-normalizer pywin32

Pack (optional):
    pyinstaller -F -w -n “一键全角空格缩进工具” Full_Width_Formatter.py

"""
from __future__ import annotations
import os
import sys
import logging
import json
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional
from datetime import datetime

# --- Third-party ---
from charset_normalizer import from_bytes
from docx import Document
from docx.text.paragraph import Paragraph

from PySide6 import QtCore, QtWidgets, QtGui

FULL_WIDTH_SPACE = "\u3000"  # U+3000
FW2 = FULL_WIDTH_SPACE * 2

# -----------------------------
# Localization
# -----------------------------

TEXTS = {
    # Window titles
    "app_title": "一键全角空格缩进",
    "output_dir_title": "选择输出文件夹",
    "clipboard_title": "剪贴板模式 - 一键全角空格缩进",
    
    # Main UI
    "main_title": "批量首行缩进（仅支持txt/docx格式）",
    "note1": "请将 .doc 文件先转为 .docx 文件，才能使用本脚本！",
    "note2": "使用前建议先清除原有排版（仅保留分段）",
    "btn_pick_files": "选择文件（可多选）…",
    "btn_clear": "清除",
    "btn_clipboard_mode": "剪贴板模式",
    "no_files_selected": "未选择文件",
    "output_placeholder": "输出文件夹路径（可直接输入或右侧选择）",
    "btn_output_dir": "输出地址…",
    "btn_start": "开始",
    "status_group": "状态",
    "processing_label": "正在处理：－",
    "done_label": "已完成：－",
    "error_label": "错误：－",
    "btn_open_output": "打开输出文件夹",
    
    # Clipboard mode UI
    "input_placeholder": "在此粘贴或输入要处理的文本...\n\n快捷键提示：\n• Ctrl+Z：撤销\n• Ctrl+C：复制\n• Ctrl+X：剪切\n• Ctrl+V：粘贴",
    "result_placeholder": "处理结果将显示在这里...",
    "btn_process": "处理",
    "btn_clear_input": "清空",
    "btn_select_all_copy": "全选复制",
    
    # Checkboxes
    "chk_clear_list": "清除转换列表",
    "chk_multipath": "多路径文件",
    "chk_enable_log": "记录日志",
    
    # Output path behavior radio buttons
    "radio_clear_on_exit": "路径仅本次使用（默认）",
    "radio_fixed_path": "路径长期固定",
    "radio_clear_after_conversion": "路径转换后清除",
    
    # File dialog
    "select_files": "选择文件",
    "file_filter": "文本与 Word (*.txt *.docx)",
    
    # Output directory dialog
    "current_path": "当前路径:",
    "btn_up": "上级",
    "btn_select_folder": "选择文件夹",
    "btn_cancel": "取消",
    "quick_access_select": "选择地址…",
    "quick_access_desktop": "🖥️ 桌面",
    "quick_access_drive": "🗃️ {0}盘",
    "quick_access_section": "快速访问",
    "folder_icon": "📁 ..",
    "cannot_access": "❌ 无法访问此目录",
    "error_prefix": "❌ 错误: {0}",
    "empty_directory": "（当前路径没有子文件夹与指定格式文件！）",
    
    # Status messages
    "processing_prefix": "正在处理：",
    "processing_complete": "正在处理：完成",
    "processing_all_complete": "正在处理：全部完成",
    "done_prefix": "已完成：",
    "error_prefix_status": "错误：",
    "more_files": "\n……等 {0} 个",
    
    # Message boxes
    "info_title": "提示",
    "success_title": "成功",
    "error_title": "错误",
    "permission_error_title": "权限错误",
    "create_failed_title": "创建失败",
    "complete_title": "完成",
    
    "msg_select_output_first": "请先选择输出文件夹",
    "msg_output_not_exist": "输出文件夹不存在",
    "msg_cannot_open_folder": "无法打开文件夹：{0}",
    "msg_select_files_first": "请先选择文件（.txt / .docx）",
    "msg_select_output_folder": "请先选择或输入输出文件夹",
    "msg_parent_not_exist": "父级目录不存在：{0}",
    "msg_no_permission_parent": "没有在父级目录创建文件夹的权限：{0}",
    "msg_folder_created": "已创建输出文件夹：{0}",
    "msg_no_permission_create": "没有权限创建文件夹：{0}",
    "msg_cannot_create_folder": "无法创建文件夹：{0}\n错误：{1}",
    "msg_unknown_error_create": "创建文件夹时发生未知错误：{0}",
    "msg_no_write_permission": "没有输出文件夹的写入权限",
    "msg_processing_complete": "处理完毕！",
    "msg_log_warning": "日志功能初始化失败，可能是权限问题。\n是否继续不记录日志？",
    "msg_copied_to_clipboard": "已复制到剪贴板",
    
    # Log messages
    "log_conversion_start": "转换开始",
    "log_conversion_list": "转换列表：",
    "log_output_dir": "输出目录: {0}",
    "log_output_list": "输出列表：",
    "log_error_list": "错误记录：",
    "log_conversion_complete": "转换完成 - 成功: {0}, 失败: {1}",
    "log_warning_title": "日志警告",
}

# -----------------------------
# Text processing utilities
# -----------------------------

def process_text_lines(text: str) -> str:
    """Process text lines by adding full-width spaces at the beginning of each non-empty line"""
    if not text:
        return text
    
    lines = text.splitlines(keepends=True)
    
    def transform_line(line: str) -> str:
        # Separate content and line ending
        if line.endswith("\r\n"):
            core, eol = line[:-2], "\r\n"
        elif line.endswith("\n"):
            core, eol = line[:-1], "\n"
        elif line.endswith("\r"):
            core, eol = line[:-1], "\r"
        else:
            core, eol = line, ""

        if core.strip() == "":
            return core + eol
        # Idempotent: ensure startswith 2 full-width spaces
        if core.startswith(FW2):
            return core + eol
        # If there are ASCII spaces/tabs at start, normalize to two full-width spaces
        core_no_leading = core.lstrip(" \t")
        return FW2 + core_no_leading + eol

    return "".join(transform_line(l) for l in lines)

# -----------------------------
# Clipboard Mode Window
# -----------------------------

class ClipboardModeWindow(QtWidgets.QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle(TEXTS["clipboard_title"])
        self.setMinimumSize(1000, 700)
        self.resize(1200, 800)
        
        # History for undo functionality
        self.input_history = []
        self.history_index = -1
        
        self._setup_ui()
        self._setup_shortcuts()
        self._apply_style()
        
    def _setup_ui(self):
        central_widget = QtWidgets.QWidget()
        self.setCentralWidget(central_widget)
        
        layout = QtWidgets.QHBoxLayout(central_widget)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(15)
        
        # Left side - Result area
        left_layout = QtWidgets.QVBoxLayout()
        
        result_label = QtWidgets.QLabel("处理结果")
        result_label.setStyleSheet("font-size: 16px; font-weight: 600; color: #333; margin-bottom: 5px;")
        left_layout.addWidget(result_label)
        
        self.result_text = QtWidgets.QTextEdit()
        self.result_text.setPlaceholderText(TEXTS["result_placeholder"])
        self.result_text.setReadOnly(True)
        left_layout.addWidget(self.result_text)
        
        # Result buttons
        result_btn_layout = QtWidgets.QHBoxLayout()
        result_btn_layout.addStretch()
        
        self.btn_select_all_copy = QtWidgets.QPushButton(TEXTS["btn_select_all_copy"])
        self.btn_select_all_copy.clicked.connect(self.select_all_copy)
        result_btn_layout.addWidget(self.btn_select_all_copy)
        
        left_layout.addLayout(result_btn_layout)
        
        # Right side - Input area
        right_layout = QtWidgets.QVBoxLayout()
        
        input_label = QtWidgets.QLabel("输入区域")
        input_label.setStyleSheet("font-size: 16px; font-weight: 600; color: #333; margin-bottom: 5px;")
        right_layout.addWidget(input_label)
        
        self.input_text = QtWidgets.QTextEdit()
        self.input_text.setPlaceholderText(TEXTS["input_placeholder"])
        self.input_text.textChanged.connect(self.on_input_changed)
        right_layout.addWidget(self.input_text)
        
        # Input buttons
        input_btn_layout = QtWidgets.QHBoxLayout()
        
        self.btn_clear_input = QtWidgets.QPushButton(TEXTS["btn_clear_input"])
        self.btn_clear_input.clicked.connect(self.clear_input)
        input_btn_layout.addWidget(self.btn_clear_input)
        
        input_btn_layout.addStretch()
        
        self.btn_process = QtWidgets.QPushButton(TEXTS["btn_process"])
        self.btn_process.setStyleSheet("font-weight: 600; min-width: 100px;")
        self.btn_process.clicked.connect(self.process_text)
        input_btn_layout.addWidget(self.btn_process)
        
        right_layout.addLayout(input_btn_layout)
        
        # Add layouts to main layout
        layout.addLayout(left_layout, 1)
        layout.addLayout(right_layout, 1)
        
    def _setup_shortcuts(self):
        # Ctrl+Z for undo (custom implementation)
        undo_shortcut = QtGui.QShortcut(QtGui.QKeySequence.Undo, self.input_text)
        undo_shortcut.activated.connect(self.custom_undo)
        
    def keyPressEvent(self, event):
        """Handle key press events for the main window"""
        # Let parent handle all key events
        super().keyPressEvent(event)
        
    def _apply_style(self):
        self.setStyleSheet("""
            QMainWindow { background: white; }
            QTextEdit {
                border: 1px solid #ddd;
                border-radius: 8px;
                padding: 10px;
                font-size: 14px;
                font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
                line-height: 1.4;
            }
            QTextEdit:focus {
                border: 2px solid #6aa3ff;
            }
            QPushButton {
                padding: 8px 16px;
                border: 1px solid #ddd;
                border-radius: 8px;
                background: #fafafa;
                font-size: 14px;
                min-height: 20px;
            }
            QPushButton:hover {
                background: #f2f2f2;
            }
            QPushButton:pressed {
                background: #e8e8e8;
            }
            QLabel {
                color: #333;
            }
        """)
        
    def on_input_changed(self):
        # Save to history for undo
        current_text = self.input_text.toPlainText()
        
        # Only save if text is different from last entry
        if not self.input_history or self.input_history[-1] != current_text:
            self.input_history.append(current_text)
            # Keep only last 50 states
            if len(self.input_history) > 50:
                self.input_history.pop(0)
            self.history_index = len(self.input_history) - 1
            
    def custom_undo(self):
        if len(self.input_history) > 1 and self.history_index > 0:
            self.history_index -= 1
            previous_text = self.input_history[self.history_index]
            
            # Temporarily disconnect signal to avoid adding to history
            self.input_text.textChanged.disconnect()
            self.input_text.setPlainText(previous_text)
            self.input_text.textChanged.connect(self.on_input_changed)
            
    def process_text(self):
        input_text = self.input_text.toPlainText()
        if not input_text.strip():
            return
            
        processed_text = process_text_lines(input_text)
        self.result_text.setPlainText(processed_text)
        
    def clear_input(self):
        self.input_text.clear()
        self.result_text.clear()
        # Reset history
        self.input_history = [""]
        self.history_index = 0
        
    def select_all_copy(self):
        result_text = self.result_text.toPlainText()
        if result_text:
            clipboard = QtWidgets.QApplication.clipboard()
            clipboard.setText(result_text)
            
            # Show temporary message
            self.show_status_message(TEXTS["msg_copied_to_clipboard"])
            
    def show_status_message(self, message):
        # Create a temporary status bar message
        if hasattr(self, 'status_bar'):
            self.status_bar.deleteLater()
            
        self.status_bar = QtWidgets.QLabel(message)
        self.status_bar.setStyleSheet("""
            QLabel {
                background: #4CAF50;
                color: white;
                padding: 8px 15px;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        self.status_bar.setParent(self)
        self.status_bar.move(10, self.height() - 40)
        self.status_bar.show()
        
        # Auto-hide after 2 seconds
        QtCore.QTimer.singleShot(2000, lambda: self.status_bar.deleteLater() if hasattr(self, 'status_bar') else None)

# -----------------------------
# Logging setup
# -----------------------------

def setup_logger(enable_log=False):
    """设置日志记录器"""
    if not enable_log:
        return None
    
    try:
        # 获取脚本所在目录
        script_dir = Path(__file__).parent
        log_dir = script_dir / "Logs"
        
        # 创建Logs目录（如果不存在）
        log_dir.mkdir(exist_ok=True)
        
        # 使用当前日期作为日志文件名
        today = datetime.now().strftime("%Y-%m-%d")
        log_file = log_dir / f"formatter_{today}.log"
        
        # 创建logger
        logger = logging.getLogger('FormatterLogger')
        logger.setLevel(logging.INFO)
        
        # 清除之前的handlers（避免重复）
        logger.handlers.clear()
        
        # 创建文件handler
        handler = logging.FileHandler(log_file, encoding='utf-8', mode='a')
        formatter = logging.Formatter('%(asctime)s - %(message)s', 
                                    datefmt='%Y-%m-%d %H:%M')
        handler.setFormatter(formatter)
        logger.addHandler(handler)
        
        return logger
        
    except Exception as e:
        # 如果日志设置失败，返回None，不影响主功能
        return None

# -----------------------------
# Settings persistence
# -----------------------------

def get_settings_file() -> Path:
    """Get the settings file path in the script directory"""
    script_dir = Path(__file__).parent
    return script_dir / "formatter_settings.json"

def load_settings() -> dict:
    """Load settings from JSON file"""
    settings_file = get_settings_file()
    default_settings = {
        "output_path": "",
        "path_behavior": "clear_on_exit"  # clear_on_exit, fixed_path, clear_after_conversion
    }
    
    try:
        if settings_file.exists():
            with open(settings_file, 'r', encoding='utf-8') as f:
                settings = json.load(f)
                # Validate path_behavior
                if settings.get("path_behavior") not in ["clear_on_exit", "fixed_path", "clear_after_conversion"]:
                    settings["path_behavior"] = "clear_on_exit"
                # Validate output_path exists only for fixed_path mode
                if settings.get("path_behavior") == "fixed_path":
                    output_path = settings.get("output_path", "")
                    if not output_path or not Path(output_path).exists():
                        settings["output_path"] = ""
                        settings["path_behavior"] = "clear_on_exit"
                return settings
    except Exception:
        pass
    
    return default_settings

def save_settings(settings: dict):
    """Save settings to JSON file"""
    try:
        settings_file = get_settings_file()
        with open(settings_file, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
    except Exception:
        pass  # Silently fail if can't save settings

# -----------------------------
# Utility: Filename resolution
# -----------------------------

def resolve_output_path(out_dir: Path, src_name: str) -> Path:
    """Return a non-conflicting output path in out_dir based on src_name.
    Strategy:
      - If no conflict: name.ext
      - If conflict: name_indent.ext
      - If conflict: name_indent1.ext, name_indent2.ext, ...
    """
    base = Path(src_name).stem
    ext = Path(src_name).suffix

    candidate = out_dir / f"{base}{ext}"
    if not candidate.exists():
        return candidate

    candidate = out_dir / f"{base}_indent{ext}"
    if not candidate.exists():
        return candidate

    i = 1
    while True:
        candidate = out_dir / f"{base}_indent{i}{ext}"
        if not candidate.exists():
            return candidate
        i += 1

# -----------------------------
# .txt processing
# -----------------------------

def detect_encoding(data: bytes) -> str:
    """Detect encoding with BOM priority; fallback to charset-normalizer; default utf-8."""
    # BOMs
    if data.startswith(b"\xef\xbb\xbf"):
        return "utf-8-sig"
    if data.startswith((b"\xff\xfe\x00\x00", b"\x00\x00\xfe\xff")):
        # UTF-32 BOMs (check before UTF-16)
        return "utf-32" if data.startswith(b"\xff\xfe\x00\x00") else "utf-32"
    if data.startswith(b"\xff\xfe"):
        return "utf-16-le"
    if data.startswith(b"\xfe\xff"):
        return "utf-16-be"

    # charset-normalizer
    try:
        result = from_bytes(data).best()
        if result and result.encoding:
            return result.encoding
    except Exception:
        pass
    return "utf-8"


def process_txt_file(src: Path, out_dir: Path) -> Path:
    data = src.read_bytes()
    enc = detect_encoding(data)

    # Preserve per-line endings by operating on bytes → decode → splitlines(keepends=True)
    text = data.decode(enc, errors="replace")
    lines = text.splitlines(keepends=True)

    def transform_line(line: str) -> str:
        # Separate content and line ending
        if line.endswith("\r\n"):
            core, eol = line[:-2], "\r\n"
        elif line.endswith("\n"):
            core, eol = line[:-1], "\n"
        elif line.endswith("\r"):
            core, eol = line[:-1], "\r"
        else:
            core, eol = line, ""

        if core.strip() == "":
            return core + eol
        # Idempotent: ensure startswith 2 full-width spaces
        if core.startswith(FW2):
            return core + eol
        # If there are ASCII spaces/tabs at start, normalize to two full-width spaces
        core_no_leading = core.lstrip(" \t")
        return FW2 + core_no_leading + eol

    new_text = "".join(transform_line(l) for l in lines)

    out_path = resolve_output_path(out_dir, src.name)
    out_path.write_text(new_text, encoding=enc, errors="replace", newline="")
    return out_path

# -----------------------------
# .docx processing
# -----------------------------

HEADING_KEYWORDS = {"Heading", "标题", "Заголовок", "Überschrift", "Rubrik"}  # common locales


def is_heading_style(p: Paragraph) -> bool:
    try:
        name = p.style.name or ""
    except Exception:
        name = ""
    return any(k in name for k in HEADING_KEYWORDS)


def paragraph_has_text(p: Paragraph) -> bool:
    return (p.text or "").strip() != ""


def ensure_fw2_at_paragraph_start(p: Paragraph) -> None:
    """Prepend two full-width spaces to paragraph *content* if not present.
    - Skips empty paragraphs
    - Skips heading styles
    - Inserts after numbering/bullet (since numbering is not part of runs)
    - Idempotent: do nothing if already startswith FW2
    - Try to preserve existing runs by editing only the first run
    - Handle soft line breaks by transforming each line within run texts
    """
    if not paragraph_has_text(p):
        return
    if is_heading_style(p):
        return

    # Quick idempotency check on full paragraph text
    full = p.text
    if full.startswith(FW2):
        return

    # Ensure we only modify the very start of the paragraph content.
    if p.runs:
        first_run = p.runs[0]
        first_run.text = FW2 + first_run.text.lstrip(" \t")
    else:
        # Paragraph has no runs (rare). Set text directly.
        p.add_run(FW2)

    # Soft line breaks inside runs: add FW2 after each break if non-empty follows
    for run in p.runs:
        if "\n" in run.text:
            parts = run.text.split("\n")
            for i in range(1, len(parts)):
                if parts[i] and not parts[i].startswith(FW2):
                    parts[i] = FW2 + parts[i].lstrip(" \t")
            run.text = "\n".join(parts)


def process_docx_file(src: Path, out_dir: Path) -> Path:
    doc = Document(str(src))
    for p in doc.paragraphs:
        ensure_fw2_at_paragraph_start(p)
    # Tables: also iterate cells' paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    ensure_fw2_at_paragraph_start(p)
    out_path = resolve_output_path(out_dir, src.name)
    doc.save(str(out_path))
    return out_path

# -----------------------------
# Worker thread
# -----------------------------

@dataclass
class TaskResult:
    src: Path
    dst: Optional[Path]
    ok: bool
    message: str


class ProcessorWorker(QtCore.QThread):
    progress = QtCore.Signal(object)  # TaskResult per file
    finished_all = QtCore.Signal()

    def __init__(self, files: List[Path], out_dir: Path, logger=None):
        super().__init__()
        self.files = files
        self.out_dir = out_dir
        self.logger = logger
        self.successful_outputs = []
        self.errors = []

    def run(self):
        # 记录转换开始
        if self.logger:
            self.logger.info("=" * 50)
            self.logger.info(TEXTS["log_conversion_start"])
            self.logger.info(TEXTS["log_conversion_list"])
            for file_path in self.files:
                self.logger.info(f"  {file_path}")
            self.logger.info(TEXTS["log_output_dir"].format(self.out_dir))
        
        for f in self.files:
            try:
                ext = f.suffix.lower()
                if ext == ".txt":
                    dst = process_txt_file(f, self.out_dir)
                elif ext == ".docx":
                    dst = process_docx_file(f, self.out_dir)
                else:
                    raise ValueError("Unsupported extension")
                
                self.successful_outputs.append(dst)
                self.progress.emit(TaskResult(f, dst, True, "Done"))
                
            except Exception as e:
                error_msg = f"{f.name}: {str(e)}"
                self.errors.append(error_msg)
                self.progress.emit(TaskResult(f, None, False, str(e)))
        
        # 记录最终结果
        if self.logger:
            self.logger.info(TEXTS["log_output_list"])
            for output_path in self.successful_outputs:
                self.logger.info(f"  {output_path}")
            
            if self.errors:
                self.logger.info(TEXTS["log_error_list"])
                for error in self.errors:
                    self.logger.info(f"  {error}")
            
            success_count = len(self.successful_outputs)
            error_count = len(self.errors)
            self.logger.info(TEXTS["log_conversion_complete"].format(success_count, error_count))
        
        self.finished_all.emit()

# -----------------------------
# Custom Output Directory Dialog
# -----------------------------

class OutputDirDialog(QtWidgets.QDialog):
    def __init__(self, parent=None, initial_path=""):
        super().__init__(parent)
        self.setWindowTitle(TEXTS["output_dir_title"])
        self.setModal(True)
        self.resize(750, 500)  # Increased size for better usability
        self.selected_path = ""
        
        # Set initial directory
        self.current_path = Path(initial_path) if initial_path and Path(initial_path).exists() else Path.cwd()
        
        self._setup_ui()
        self._apply_style()
        self._load_directory()
        
    def _setup_ui(self):
        layout = QtWidgets.QVBoxLayout(self)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(10)
        
        # Path bar with dropdown
        path_layout = QtWidgets.QHBoxLayout()
        
        # Quick access dropdown
        self.quick_combo = QtWidgets.QComboBox()
        self.quick_combo.setMinimumWidth(100)  # Reduced from 120
        self.quick_combo.setMaximumWidth(140)  # Set maximum width to prevent it from expanding too much
        self._populate_quick_access()
        self.quick_combo.currentTextChanged.connect(self._on_quick_access_changed)
        path_layout.addWidget(self.quick_combo)
        
        path_layout.addWidget(QtWidgets.QLabel(TEXTS["current_path"]))
        
        self.path_edit = QtWidgets.QLineEdit()
        self.path_edit.setText(str(self.current_path))
        self.path_edit.returnPressed.connect(self._navigate_to_path)
        path_layout.addWidget(self.path_edit, 1)
        
        self.btn_up = QtWidgets.QPushButton(TEXTS["btn_up"])
        self.btn_up.clicked.connect(self._go_up)
        path_layout.addWidget(self.btn_up)
        
        layout.addLayout(path_layout)
        
        # File list
        self.file_list = QtWidgets.QListWidget()
        self.file_list.itemDoubleClicked.connect(self._item_double_clicked)
        layout.addWidget(self.file_list)
        
        # Button layout
        button_layout = QtWidgets.QHBoxLayout()
        button_layout.addStretch()
        
        self.btn_select = QtWidgets.QPushButton(TEXTS["btn_select_folder"])
        self.btn_select.clicked.connect(self._select_current)
        button_layout.addWidget(self.btn_select)
        
        self.btn_cancel = QtWidgets.QPushButton(TEXTS["btn_cancel"])
        self.btn_cancel.clicked.connect(self.reject)
        button_layout.addWidget(self.btn_cancel)
        
        layout.addLayout(button_layout)
        
    def _populate_quick_access(self):
        """Populate quick access dropdown with drives and special folders"""
        self.quick_combo.addItem(TEXTS["quick_access_select"], "")
        
        # Add desktop
        try:
            desktop_path = Path.home() / "Desktop"
            if desktop_path.exists():
                self.quick_combo.addItem(TEXTS["quick_access_desktop"], str(desktop_path))
        except:
            pass
        
        # Add available drives
        import string
        for drive_letter in string.ascii_uppercase:
            drive_path = Path(f"{drive_letter}:\\")
            if drive_path.exists():
                self.quick_combo.addItem(TEXTS["quick_access_drive"].format(drive_letter), str(drive_path))
        
        # Add Windows Quick Access using pywin32
        quick_access_items = []
        try:
            import win32com.client
            shell = win32com.client.Dispatch("Shell.Application")
            # Get Quick Access folder (FOLDERID_QuickAccess)
            quick_access = shell.Namespace("shell:::{679f85cb-0220-4080-b29b-5540cc05aab6}")
            if quick_access:
                for item in quick_access.Items():
                    try:
                        path_str = item.Path
                        if path_str and Path(path_str).exists() and Path(path_str).is_dir():
                            # Skip if it's already added (like Desktop)
                            already_added = False
                            for i in range(1, self.quick_combo.count()):
                                if self.quick_combo.itemData(i) == path_str:
                                    already_added = True
                                    break
                            if not already_added:
                                display_name = item.Name or Path(path_str).name
                                quick_access_items.append((f"↘️ {display_name}", path_str))
                    except:
                        continue
        except ImportError:
            # pywin32 not available, skip Quick Access
            pass
        except Exception:
            # Other errors, skip Quick Access
            pass
        
        # Add separator and quick access items if any exist
        if quick_access_items:
            # Add separator (non-clickable)
            self.quick_combo.addItem(TEXTS["quick_access_section"], "")
            # Disable the separator item
            separator_index = self.quick_combo.count() - 1
            separator_item = self.quick_combo.model().item(separator_index)
            separator_item.setEnabled(False)
            
            # Add the actual quick access items
            for display_name, path_str in quick_access_items:
                self.quick_combo.addItem(display_name, path_str)
        
    def _on_quick_access_changed(self, text):
        """Handle quick access selection"""
        if text == TEXTS["quick_access_select"] or text == TEXTS["quick_access_section"]:
            return
            
        # Get the path from combo data
        current_index = self.quick_combo.currentIndex()
        if (current_index > 0):  # Skip the first TEXTS["quick_access_select"] item
            path_str = self.quick_combo.itemData(current_index)
            if path_str:
                try:
                    new_path = Path(path_str)
                    if new_path.exists() and new_path.is_dir():
                        self.current_path = new_path
                        self._load_directory()
                except Exception:
                    pass
        
        # Reset combo to TEXTS["quick_access_select"]
        self.quick_combo.setCurrentIndex(0)
        
    def _apply_style(self):
        self.setStyleSheet("""
            QDialog { background: white; }
            QLineEdit { 
                padding: 8px 10px; 
                border: 1px solid #ddd; 
                border-radius: 8px; 
                font-size: 14px;
            }
            QComboBox {
                padding: 6px 8px;
                border: 1px solid #ddd;
                border-radius: 8px;
                background: #fafafa;
                font-size: 14px;
            }
            QComboBox:hover {
                background: #f2f2f2;
            }
            QComboBox::drop-down {
                border: none;
                padding-right: 8px;
            }
            QComboBox::down-arrow {
                image: none;
                border: 2px solid #666;
                border-top: none;
                border-left: none;
                width: 6px;
                height: 6px;
                transform: rotate(45deg);
                margin-top: -3px;
            }
            QComboBox QAbstractItemView::item:disabled {
                color: #999;
                background: #f8f8f8;
            }
            QComboBox QAbstractItemView::item:disabled:hover {
                background: #f8f8f8;
            }
            QPushButton { 
                padding: 8px 15px; 
                border: 1px solid #ddd; 
                border-radius: 8px; 
                background: #fafafa; 
                min-width: 70px;
                font-size: 14px;
            }
            QPushButton:hover { background: #f2f2f2; }
            QListWidget { 
                border: 1px solid #ddd; 
                border-radius: 8px; 
                background: white;
                font-size: 14px;
            }
            QListWidget::item {
                padding: 6px 10px;
                border-bottom: 1px solid #f0f0f0;
            }
            QListWidget::item:hover {
                background: #f5f5f5;
            }
            QListWidget::item:selected {
                background: #e3f2fd;
                color: black;
            }
            QLabel {
                font-size: 14px;
            }
        """)
        
    def _load_directory(self):
        self.file_list.clear()
        
        try:
            if not self.current_path.exists():
                self.current_path = Path.cwd()
                
            self.path_edit.setText(str(self.current_path))
            
            # Add parent directory item (if not root)
            if self.current_path.parent != self.current_path:
                item = QtWidgets.QListWidgetItem(TEXTS["folder_icon"])
                item.setData(QtCore.Qt.UserRole, str(self.current_path.parent))
                self.file_list.addItem(item)
            
            # Get all items in directory
            items = []
            try:
                for path in self.current_path.iterdir():
                    if path.is_dir():
                        items.append((f"📁 {path.name}", str(path), True))
                    elif path.suffix.lower() in {'.txt', '.docx'}:
                        icon = "📄" if path.suffix.lower() == '.txt' else "📝"
                        items.append((f"{icon} {path.name}", str(path), False))
            except PermissionError:
                item = QtWidgets.QListWidgetItem(TEXTS["cannot_access"])
                self.file_list.addItem(item)
                return
                
            # Sort: directories first, then files
            items.sort(key=lambda x: (not x[2], x[0].lower()))
            
            # Add items to list
            for display_name, full_path, is_dir in items:
                item = QtWidgets.QListWidgetItem(display_name)
                item.setData(QtCore.Qt.UserRole, full_path)
                self.file_list.addItem(item)
                
            # Show empty directory message if no items
            if not items:
                item = QtWidgets.QListWidgetItem(TEXTS["empty_directory"])
                self.file_list.addItem(item)
                
        except Exception as e:
            item = QtWidgets.QListWidgetItem(TEXTS["error_prefix"].format(str(e)))
            self.file_list.addItem(item)
            
    def _navigate_to_path(self):
        path_text = self.path_edit.text().strip()
        try:
            new_path = Path(path_text)
            if new_path.exists() and new_path.is_dir():
                self.current_path = new_path
                self._load_directory()
        except Exception:
            # Reset to current path if invalid
            self.path_edit.setText(str(self.current_path))
            
    def _go_up(self):
        if self.current_path.parent != self.current_path:
            self.current_path = self.current_path.parent
            self._load_directory()
            
    def _item_double_clicked(self, item):
        path_str = item.data(QtCore.Qt.UserRole)
        if path_str:
            path = Path(path_str)
            if path.exists() and path.is_dir():
                self.current_path = path
                self._load_directory()
                
    def _select_current(self):
        # Use the path from the text field instead of current_path
        # This allows users to type non-existent paths that will be created later
        typed_path = self.path_edit.text().strip()
        if typed_path:
            self.selected_path = typed_path
        else:
            self.selected_path = str(self.current_path)
        self.accept()
        
    def get_selected_path(self):
        return self.selected_path

# -----------------------------
# GUI
# -----------------------------

class IndentorApp(QtWidgets.QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(TEXTS["app_title"])
        self.setMinimumWidth(800)  # Increased width to accommodate longer text
        self.setMinimumHeight(580)  # Further reduced from 600 to 580 due to smaller margins
        self.files: List[Path] = []
        self.out_dir: Optional[Path] = None
        self.worker: Optional[ProcessorWorker] = None
        self.converted_outputs: List[Path] = []  # Track converted files
        self.settings = load_settings()  # Load settings
        self._build_ui()
        self._apply_style()
        self._load_settings()  # Apply settings

    def _build_ui(self):
        layout = QtWidgets.QVBoxLayout(self)
        layout.setContentsMargins(18, 15, 18, 18)  # Slightly increased top margin for better balance
        layout.setSpacing(10)  # Increased from 8 to 10 for better spacing

        # Title and clipboard button row
        title_row = QtWidgets.QHBoxLayout()
        
        title = QtWidgets.QLabel(TEXTS["main_title"])
        title.setStyleSheet("font-size:20px; font-weight:600;")
        title_row.addWidget(title)
        
        title_row.addStretch()  # Push clipboard button to the right
        
        # Clipboard mode button (moved to top right with increased size)
        self.btn_clipboard_mode = QtWidgets.QPushButton(TEXTS["btn_clipboard_mode"])
        self.btn_clipboard_mode.setFixedHeight(40)  # Increased from 32 to 40
        self.btn_clipboard_mode.setFixedWidth(120)  # Increased from 100 to 120
        self.btn_clipboard_mode.setStyleSheet("font-weight:600; font-size:13px;")  # Increased font size
        self.btn_clipboard_mode.clicked.connect(self.open_clipboard_mode)
        title_row.addWidget(self.btn_clipboard_mode)
        
        layout.addLayout(title_row)

        # Add spacing after title
        layout.addSpacing(5)

        # Notes with optimized spacing
        notes_layout = QtWidgets.QVBoxLayout()
        notes_layout.setSpacing(3)  # Slightly increased from 2 to 3
        
        note1 = QtWidgets.QLabel(TEXTS["note1"])
        note1.setStyleSheet("color:#666;")
        note1.setAlignment(QtCore.Qt.AlignLeft)
        notes_layout.addWidget(note1)
        
        note2 = QtWidgets.QLabel(TEXTS["note2"])
        note2.setStyleSheet("color:#666;")
        note2.setAlignment(QtCore.Qt.AlignLeft)
        notes_layout.addWidget(note2)
        
        layout.addLayout(notes_layout)

        # Add spacing after notes
        layout.addSpacing(8)

        # Multi-path checkbox row with better spacing
        checkbox_row = QtWidgets.QHBoxLayout()
        
        self.chk_clear_list = QtWidgets.QCheckBox(TEXTS["chk_clear_list"])
        self.chk_clear_list.setChecked(True)  # Default checked
        self.chk_clear_list.setStyleSheet("color:#666;")
        checkbox_row.addWidget(self.chk_clear_list)
        
        # Add spacing between checkboxes
        checkbox_row.addSpacing(25)  # Increased from 20 to 25
        
        self.chk_multipath = QtWidgets.QCheckBox(TEXTS["chk_multipath"])
        self.chk_multipath.setStyleSheet("color:#666;")
        checkbox_row.addWidget(self.chk_multipath)
        
        # Add spacing and log checkbox
        checkbox_row.addSpacing(25)  # Increased from 20 to 25
        self.chk_enable_log = QtWidgets.QCheckBox(TEXTS["chk_enable_log"])
        self.chk_enable_log.setStyleSheet("color:#666;")
        checkbox_row.addWidget(self.chk_enable_log)
        
        checkbox_row.addStretch()  # Push checkboxes to the left
        layout.addLayout(checkbox_row)

        # Add spacing before file picker row
        layout.addSpacing(6)  # Increased from 2 to 6

        # File picker row
        file_row = QtWidgets.QHBoxLayout()
        self.btn_pick = QtWidgets.QPushButton(TEXTS["btn_pick_files"])
        self.btn_pick.clicked.connect(self.pick_files)
        file_row.addWidget(self.btn_pick)
        
        self.btn_clear = QtWidgets.QPushButton(TEXTS["btn_clear"])
        self.btn_clear.clicked.connect(self.clear_files)
        file_row.addWidget(self.btn_clear)

        self.sel_summary = QtWidgets.QLabel(TEXTS["no_files_selected"])
        self.sel_summary.setStyleSheet("color:#444;")
        self.sel_summary.setTextInteractionFlags(QtCore.Qt.TextSelectableByMouse)
        file_row.addWidget(self.sel_summary, 1)
        layout.addLayout(file_row)

        # Add spacing before output dir row
        layout.addSpacing(4)

        # Output dir row
        out_row = QtWidgets.QHBoxLayout()
        self.out_edit = QtWidgets.QLineEdit()
        self.out_edit.setPlaceholderText(TEXTS["output_placeholder"])
        out_row.addWidget(self.out_edit, 1)
        self.btn_out = QtWidgets.QPushButton(TEXTS["btn_output_dir"])
        self.btn_out.clicked.connect(self.pick_out_dir)
        out_row.addWidget(self.btn_out)
        layout.addLayout(out_row)

        # Add spacing before radio buttons
        layout.addSpacing(6)  # Increased from 2 to 6

        # Output path behavior radio buttons
        radio_row = QtWidgets.QHBoxLayout()
        self.radio_clear_on_exit = QtWidgets.QRadioButton(TEXTS["radio_clear_on_exit"])
        self.radio_clear_on_exit.setChecked(True)  # Default checked
        radio_row.addWidget(self.radio_clear_on_exit)
        
        self.radio_fixed_path = QtWidgets.QRadioButton(TEXTS["radio_fixed_path"])
        radio_row.addWidget(self.radio_fixed_path)
        
        self.radio_clear_after_conversion = QtWidgets.QRadioButton(TEXTS["radio_clear_after_conversion"])
        radio_row.addWidget(self.radio_clear_after_conversion)
        
        radio_row.addStretch()  # Push radio buttons to the left
        layout.addLayout(radio_row)

        # Add spacing before start button
        layout.addSpacing(8)  # Increased from 2 to 8

        # Start button
        self.btn_start = QtWidgets.QPushButton(TEXTS["btn_start"])
        self.btn_start.setFixedHeight(36)
        self.btn_start.setStyleSheet("font-weight:600;")
        self.btn_start.clicked.connect(self.start_processing)
        layout.addWidget(self.btn_start)

        # Add spacing before status area
        layout.addSpacing(5)

        # Status area
        grp = QtWidgets.QGroupBox(TEXTS["status_group"])
        gl = QtWidgets.QGridLayout(grp)
        gl.setContentsMargins(12, 12, 12, 12)
        gl.setHorizontalSpacing(10)
        gl.setVerticalSpacing(6)

        self.lbl_processing = QtWidgets.QLabel(TEXTS["processing_label"])
        self.lbl_done = QtWidgets.QLabel(TEXTS["done_label"])
        self.lbl_err = QtWidgets.QLabel(TEXTS["error_label"])
        for w in (self.lbl_processing, self.lbl_done, self.lbl_err):
            w.setTextInteractionFlags(QtCore.Qt.TextSelectableByMouse)

        gl.addWidget(self.lbl_processing, 0, 0)
        gl.addWidget(self.lbl_done, 1, 0)
        gl.addWidget(self.lbl_err, 2, 0)

        self.btn_open_out = QtWidgets.QPushButton(TEXTS["btn_open_output"])
        self.btn_open_out.clicked.connect(self.open_out_dir)
        gl.addWidget(self.btn_open_out, 3, 0, 1, 1)

        layout.addWidget(grp)

        # Add minimal spacing before progress bar
        layout.addSpacing(3)

        self.progress_bar = QtWidgets.QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar)

    def _apply_style(self):
        # Minimal, modern, light theme
        self.setStyleSheet(
            """
            QWidget { font-size:14px; }
            QPushButton { padding:6px 12px; border:1px solid #ddd; border-radius:8px; background:#fafafa; }
            QPushButton:hover { background:#f2f2f2; }
            QLineEdit { padding:6px 8px; border:1px solid #ddd; border-radius:8px; }
            QGroupBox { border:1px solid #eee; border-radius:10px; margin-top:10px; }
            QGroupBox::title { subcontrol-origin: margin; left:10px; padding:0 4px; color:#666; }
            QProgressBar { height:12px; border:1px solid #ddd; border-radius:6px; background:#f5f5f5; }
            QProgressBar::chunk { border-radius:6px; background:#6aa3ff; }
            """
        )

    def _load_settings(self):
        """Apply loaded settings to the UI"""
        self.out_edit.setText(self.settings.get("output_path", ""))
        path_behavior = self.settings.get("path_behavior", "clear_on_exit")
        if path_behavior == "fixed_path":
            self.radio_fixed_path.setChecked(True)
        elif path_behavior == "clear_after_conversion":
            self.radio_clear_after_conversion.setChecked(True)
        else:
            self.radio_clear_on_exit.setChecked(True)

    def open_clipboard_mode(self):
        self.clipboard_window = ClipboardModeWindow(self)
        self.clipboard_window.show()

    def pick_files(self):
        options = QtWidgets.QFileDialog.Options()
        files, _ = QtWidgets.QFileDialog.getOpenFileNames(
            self, TEXTS["select_files"], "", TEXTS["file_filter"], options=options
        )
        if files:
            if self.chk_clear_list.isChecked():
                self.files = []
            self.files.extend(Path(f) for f in files)
            self.update_file_summary()

    def clear_files(self):
        self.files = []
        self.update_file_summary()

    def update_file_summary(self):
        if not self.files:
            self.sel_summary.setText(TEXTS["no_files_selected"])
        else:
            file_names = [f.name for f in self.files[:3]]
            more_count = len(self.files) - 3
            if more_count > 0:
                file_names.append(TEXTS["more_files"].format(more_count))
            self.sel_summary.setText("\n".join(file_names))

    def pick_out_dir(self):
        dialog = OutputDirDialog(self, initial_path=self.out_edit.text())
        if dialog.exec() == QtWidgets.QDialog.Accepted:
            self.out_dir = Path(dialog.get_selected_path())
            self.out_edit.setText(str(self.out_dir))

    def start_processing(self):
        if not self.files:
            QtWidgets.QMessageBox.warning(self, TEXTS["error_title"], TEXTS["msg_select_files_first"])
            return

        out_dir_text = self.out_edit.text().strip()
        if not out_dir_text:
            QtWidgets.QMessageBox.warning(self, TEXTS["error_title"], TEXTS["msg_select_output_folder"])
            return

        self.out_dir = Path(out_dir_text)
        if not self.out_dir.exists():
            try:
                self.out_dir.mkdir(parents=True, exist_ok=True)
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, TEXTS["create_failed_title"], str(e))
                return

        if not os.access(self.out_dir, os.W_OK):
            QtWidgets.QMessageBox.critical(self, TEXTS["permission_error_title"], TEXTS["msg_no_write_permission"])
            return

        # Save settings
        self.settings["output_path"] = str(self.out_dir)
        if self.radio_fixed_path.isChecked():
            self.settings["path_behavior"] = "fixed_path"
        elif self.radio_clear_after_conversion.isChecked():
            self.settings["path_behavior"] = "clear_after_conversion"
        else:
            self.settings["path_behavior"] = "clear_on_exit"
        save_settings(self.settings)

        # Setup logger
        enable_log = self.chk_enable_log.isChecked()
        self.logger = setup_logger(enable_log)
        if enable_log and not self.logger:
            reply = QtWidgets.QMessageBox.warning(
                self, TEXTS["log_warning_title"], TEXTS["msg_log_warning"], 
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No, QtWidgets.QMessageBox.No
            )
            if reply == QtWidgets.QMessageBox.No:
                return

        self.converted_outputs = []
        self.lbl_processing.setText(TEXTS["processing_label"])
        self.lbl_done.setText(TEXTS["done_label"])
        self.lbl_err.setText(TEXTS["error_label"])
        self.progress_bar.setValue(0)

        self.worker = ProcessorWorker(self.files, self.out_dir, self.logger)
        self.worker.progress.connect(self.on_progress)
        self.worker.finished_all.connect(self.on_finished_all)
        self.worker.start()

    def on_progress(self, result: TaskResult):
        if result.ok:
            self.converted_outputs.append(result.dst)
            self.lbl_done.setText(f"{TEXTS['done_prefix']} {len(self.converted_outputs)}")
        else:
            self.lbl_err.setText(f"{TEXTS['error_prefix_status']} {result.message}")

        progress = (len(self.converted_outputs) + len(self.worker.errors)) / len(self.files) * 100
        self.progress_bar.setValue(progress)

    def on_finished_all(self):
        self.lbl_processing.setText(TEXTS["processing_complete"])
        self.progress_bar.setValue(100)
        QtWidgets.QMessageBox.information(self, TEXTS["complete_title"], TEXTS["msg_processing_complete"])

        if self.radio_clear_after_conversion.isChecked():
            self.out_edit.clear()
            self.out_dir = None

    def open_out_dir(self):
        if self.out_dir and self.out_dir.exists():
            try:
                os.startfile(self.out_dir)
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, TEXTS["error_title"], TEXTS["msg_cannot_open_folder"].format(e))
        else:
            QtWidgets.QMessageBox.warning(self, TEXTS["error_title"], TEXTS["msg_output_not_exist"])

def main():
    app = QtWidgets.QApplication(sys.argv)
    w = IndentorApp()
    w.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
