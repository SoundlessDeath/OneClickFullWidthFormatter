r"""
Indentor v1 (Windows-only)

Batch-insert two full-width spaces (U+3000×2) at the start of each line (.txt) or each paragraph (.docx).

- Formats: .txt, .docx (NO .doc)
- Naming: if conflict → _indent; if conflict persists → _copy(2), _copy(3) ...
- Encoding (.txt): detect & preserve; preserves per-line line endings
- .docx: inserts after bullets/numbering; skips empty paragraphs; skips Heading styles
- UI: minimal, pretty; shows up to 3 selected files; shows processing & last 3 finished; has "Open output folder"
- Note line: "Please convert .doc to .docx before using this tool!"

Dependencies (install in Windows):
    pip install PySide6 python-docx charset-normalizer

Pack (optional):
    pyinstaller -F -w -n Indentor_v1_windows indentor_v1_windows.py

"""
from __future__ import annotations
import os
import sys
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

# --- Third-party ---
from charset_normalizer import from_bytes
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from PySide6 import QtCore, QtGui, QtWidgets

FULL_WIDTH_SPACE = "\u3000"  # U+3000
FW2 = FULL_WIDTH_SPACE * 2

# -----------------------------
# Utility: Filename resolution
# -----------------------------

def resolve_output_path(out_dir: Path, src_name: str) -> Path:
    """Return a non-conflicting output path in out_dir based on src_name.
    Strategy:
      - If no conflict: name.ext
      - If conflict: name_indent.ext
      - If conflict: name_indent1.ext, name_indent2.ext, ...
    """
    base = Path(src_name).stem
    ext = Path(src_name).suffix

    candidate = out_dir / f"{base}{ext}"
    if not candidate.exists():
        return candidate

    candidate = out_dir / f"{base}_indent{ext}"
    if not candidate.exists():
        return candidate

    i = 1
    while True:
        candidate = out_dir / f"{base}_indent{i}{ext}"
        if not candidate.exists():
            return candidate
        i += 1

# -----------------------------
# .txt processing
# -----------------------------

def detect_encoding(data: bytes) -> str:
    """Detect encoding with BOM priority; fallback to charset-normalizer; default utf-8."""
    # BOMs
    if data.startswith(b"\xef\xbb\xbf"):
        return "utf-8-sig"
    if data.startswith((b"\xff\xfe\x00\x00", b"\x00\x00\xfe\xff")):
        # UTF-32 BOMs (check before UTF-16)
        return "utf-32" if data.startswith(b"\xff\xfe\x00\x00") else "utf-32"
    if data.startswith(b"\xff\xfe"):
        return "utf-16-le"
    if data.startswith(b"\xfe\xff"):
        return "utf-16-be"

    # charset-normalizer
    try:
        result = from_bytes(data).best()
        if result and result.encoding:
            return result.encoding
    except Exception:
        pass
    return "utf-8"


def process_txt_file(src: Path, out_dir: Path) -> Path:
    data = src.read_bytes()
    enc = detect_encoding(data)

    # Preserve per-line endings by operating on bytes → decode → splitlines(keepends=True)
    text = data.decode(enc, errors="replace")
    lines = text.splitlines(keepends=True)

    def transform_line(line: str) -> str:
        # Separate content and line ending
        if line.endswith("\r\n"):
            core, eol = line[:-2], "\r\n"
        elif line.endswith("\n"):
            core, eol = line[:-1], "\n"
        elif line.endswith("\r"):
            core, eol = line[:-1], "\r"
        else:
            core, eol = line, ""

        if core.strip() == "":
            return core + eol
        # Idempotent: ensure startswith 2 full-width spaces
        if core.startswith(FW2):
            return core + eol
        # If there are ASCII spaces/tabs at start, normalize to two full-width spaces
        core_no_leading = core.lstrip(" \t")
        return FW2 + core_no_leading + eol

    new_text = "".join(transform_line(l) for l in lines)

    out_path = resolve_output_path(out_dir, src.name)
    out_path.write_text(new_text, encoding=enc, errors="replace", newline="")
    return out_path

# -----------------------------
# .docx processing
# -----------------------------

HEADING_KEYWORDS = {"Heading", "标题", "Заголовок", "Überschrift", "Rubrik"}  # common locales


def is_heading_style(p: Paragraph) -> bool:
    try:
        name = p.style.name or ""
    except Exception:
        name = ""
    return any(k in name for k in HEADING_KEYWORDS)


def paragraph_has_text(p: Paragraph) -> bool:
    return (p.text or "").strip() != ""


def ensure_fw2_at_paragraph_start(p: Paragraph) -> None:
    """Prepend two full-width spaces to paragraph *content* if not present.
    - Skips empty paragraphs
    - Skips heading styles
    - Inserts after numbering/bullet (since numbering is not part of runs)
    - Idempotent: do nothing if already startswith FW2
    - Try to preserve existing runs by editing only the first run
    - Handle soft line breaks by transforming each line within run texts
    """
    if not paragraph_has_text(p):
        return
    if is_heading_style(p):
        return

    # Quick idempotency check on full paragraph text
    full = p.text
    if full.startswith(FW2):
        return

    # Ensure we only modify the very start of the paragraph content.
    if p.runs:
        first_run = p.runs[0]
        first_run.text = FW2 + first_run.text.lstrip(" \t")
    else:
        # Paragraph has no runs (rare). Set text directly.
        p.add_run(FW2)

    # Soft line breaks inside runs: add FW2 after each break if non-empty follows
    for run in p.runs:
        if "\n" in run.text:
            parts = run.text.split("\n")
            for i in range(1, len(parts)):
                if parts[i] and not parts[i].startswith(FW2):
                    parts[i] = FW2 + parts[i].lstrip(" \t")
            run.text = "\n".join(parts)


def process_docx_file(src: Path, out_dir: Path) -> Path:
    doc = Document(str(src))
    for p in doc.paragraphs:
        ensure_fw2_at_paragraph_start(p)
    # Tables: also iterate cells' paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    ensure_fw2_at_paragraph_start(p)
    out_path = resolve_output_path(out_dir, src.name)
    doc.save(str(out_path))
    return out_path

# -----------------------------
# Worker thread
# -----------------------------

@dataclass
class TaskResult:
    src: Path
    dst: Optional[Path]
    ok: bool
    message: str


class ProcessorWorker(QtCore.QThread):
    progress = QtCore.Signal(object)  # TaskResult per file
    finished_all = QtCore.Signal()

    def __init__(self, files: List[Path], out_dir: Path):
        super().__init__()
        self.files = files
        self.out_dir = out_dir

    def run(self):
        for f in self.files:
            try:
                ext = f.suffix.lower()
                if ext == ".txt":
                    dst = process_txt_file(f, self.out_dir)
                elif ext == ".docx":
                    dst = process_docx_file(f, self.out_dir)
                else:
                    raise ValueError("Unsupported extension")
                self.progress.emit(TaskResult(f, dst, True, "Done"))
            except Exception as e:
                self.progress.emit(TaskResult(f, None, False, str(e)))
        self.finished_all.emit()

# -----------------------------
# GUI
# -----------------------------

class IndentorApp(QtWidgets.QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Indentor v1 · Windows")
        self.setMinimumWidth(800)  # Increased width to accommodate longer text
        self.files: List[Path] = []
        self.out_dir: Optional[Path] = None
        self.worker: Optional[ProcessorWorker] = None
        self._build_ui()
        self._apply_style()

    def _build_ui(self):
        layout = QtWidgets.QVBoxLayout(self)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(12)

        # Title
        title = QtWidgets.QLabel("批量首行缩进")
        title.setStyleSheet("font-size:20px; font-weight:600;")
        layout.addWidget(title)

        # Split note into two lines
        note1 = QtWidgets.QLabel("请将 .doc 文件先转为 .docx 文件，才能使用本脚本！")
        note1.setStyleSheet("color:#666;")
        note1.setAlignment(QtCore.Qt.AlignLeft)
        layout.addWidget(note1)
        
        note2 = QtWidgets.QLabel("使用前建议先清除原有排版（仅保留分段）")
        note2.setStyleSheet("color:#666;")
        note2.setAlignment(QtCore.Qt.AlignLeft)
        layout.addWidget(note2)

        # Multi-path checkbox
        self.chk_multipath = QtWidgets.QCheckBox("多路径文件")
        self.chk_multipath.setStyleSheet("color:#666;")
        layout.addWidget(self.chk_multipath)

        # File picker row
        file_row = QtWidgets.QHBoxLayout()
        self.btn_pick = QtWidgets.QPushButton("选择文件（可多选）…")
        self.btn_pick.clicked.connect(self.pick_files)
        file_row.addWidget(self.btn_pick)
        
        self.btn_clear = QtWidgets.QPushButton("清除")
        self.btn_clear.clicked.connect(self.clear_files)
        file_row.addWidget(self.btn_clear)

        self.sel_summary = QtWidgets.QLabel("未选择文件")
        self.sel_summary.setStyleSheet("color:#444;")
        self.sel_summary.setTextInteractionFlags(QtCore.Qt.TextSelectableByMouse)
        file_row.addWidget(self.sel_summary, 1)
        layout.addLayout(file_row)

        # Output dir row
        out_row = QtWidgets.QHBoxLayout()
        self.out_edit = QtWidgets.QLineEdit()
        self.out_edit.setPlaceholderText("输出文件夹路径（可直接输入或右侧选择）")
        out_row.addWidget(self.out_edit, 1)
        self.btn_out = QtWidgets.QPushButton("输出地址…")
        self.btn_out.clicked.connect(self.pick_out_dir)
        out_row.addWidget(self.btn_out)
        layout.addLayout(out_row)

        # Start button
        self.btn_start = QtWidgets.QPushButton("开始")
        self.btn_start.setFixedHeight(36)
        self.btn_start.setStyleSheet("font-weight:600;")
        self.btn_start.clicked.connect(self.start_processing)
        layout.addWidget(self.btn_start)

        # Status area
        grp = QtWidgets.QGroupBox("状态")
        gl = QtWidgets.QGridLayout(grp)
        gl.setContentsMargins(12, 12, 12, 12)
        gl.setHorizontalSpacing(10)
        gl.setVerticalSpacing(6)

        self.lbl_processing = QtWidgets.QLabel("正在处理：－")
        self.lbl_done = QtWidgets.QLabel("已完成：－")
        self.lbl_err = QtWidgets.QLabel("错误：－")
        for w in (self.lbl_processing, self.lbl_done, self.lbl_err):
            w.setTextInteractionFlags(QtCore.Qt.TextSelectableByMouse)

        gl.addWidget(self.lbl_processing, 0, 0)
        gl.addWidget(self.lbl_done, 1, 0)
        gl.addWidget(self.lbl_err, 2, 0)

        self.btn_open_out = QtWidgets.QPushButton("打开输出文件夹")
        self.btn_open_out.clicked.connect(self.open_out_dir)
        gl.addWidget(self.btn_open_out, 3, 0, 1, 1)

        layout.addWidget(grp)

        self.progress_bar = QtWidgets.QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar)

    def _apply_style(self):
        # Minimal, modern, light theme
        self.setStyleSheet(
            """
            QWidget { font-size:14px; }
            QPushButton { padding:6px 12px; border:1px solid #ddd; border-radius:8px; background:#fafafa; }
            QPushButton:hover { background:#f2f2f2; }
            QLineEdit { padding:6px 8px; border:1px solid #ddd; border-radius:8px; }
            QGroupBox { border:1px solid #eee; border-radius:10px; margin-top:10px; }
            QGroupBox::title { subcontrol-origin: margin; left:10px; padding:0 4px; color:#666; }
            QProgressBar { height:12px; border:1px solid #ddd; border-radius:6px; background:#f5f5f5; }
            QProgressBar::chunk { border-radius:6px; background:#6aa3ff; }
            """
        )

    # ---------------- Events ----------------
    def pick_files(self):
        dlg = QtWidgets.QFileDialog(self, "选择文件")
        dlg.setFileMode(QtWidgets.QFileDialog.ExistingFiles)
        # Only one combined filter, no "All files" option
        dlg.setNameFilters(["文本与 Word (*.txt *.docx)"])
        dlg.selectNameFilter("文本与 Word (*.txt *.docx)")
        if dlg.exec():
            selected = [Path(p) for p in dlg.selectedFiles()]
            # Filter to .txt / .docx just in case
            selected = [p for p in selected if p.suffix.lower() in {'.txt', '.docx'}]
            if self.chk_multipath.isChecked():
                self.files.extend(selected)
            else:
                self.files = selected
            self.update_selected_summary()

    def update_selected_summary(self):
        if not self.files:
            self.sel_summary.setText("未选择文件")
            return
        paths = [str(p) for p in self.files[:3]]
        extra = len(self.files) - 3
        text = "\n".join(paths)
        if extra > 0:
            text += f"\n……等 {extra} 个"
        self.sel_summary.setText(text)

    def pick_out_dir(self):
        d = QtWidgets.QFileDialog.getExistingDirectory(self, "选择输出文件夹")
        if d:
            self.out_edit.setText(d)

    def open_out_dir(self):
        path = self.out_edit.text().strip()
        if not path:
            QtWidgets.QMessageBox.information(self, "提示", "请先选择输出文件夹")
            return
        if not Path(path).exists():
            QtWidgets.QMessageBox.warning(self, "提示", "输出文件夹不存在")
            return
        # Open in Explorer with error handling
        try:
            os.startfile(path)
        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "错误", f"无法打开文件夹：{str(e)}")

    def start_processing(self):
        if not self.files:
            QtWidgets.QMessageBox.information(self, "提示", "请先选择文件（.txt / .docx）")
            return
        out_dir = self.out_edit.text().strip()
        if not out_dir:
            QtWidgets.QMessageBox.information(self, "提示", "请先选择或输入输出文件夹")
            return
        out_path = Path(out_dir)
        if not out_path.exists():
            QtWidgets.QMessageBox.information(self, "提示", "输出文件夹不存在")
            return
        if not os.access(out_path, os.W_OK):
            QtWidgets.QMessageBox.warning(self, "提示", "没有输出文件夹的写入权限")
            return

        # Clean up previous worker if exists
        if self.worker and self.worker.isRunning():
            self.worker.terminate()
            self.worker.wait()

        self.btn_start.setEnabled(False)
        self.btn_pick.setEnabled(False)
        self.btn_out.setEnabled(False)
        self.progress_bar.setValue(0)
        self.lbl_processing.setText("正在处理：－")
        self.lbl_done.setText("已完成：－")
        self.lbl_err.setText("错误：－")
        
        # Reset completion tracking
        self._completed_names = []
        self._error_msgs = []

        self.worker = ProcessorWorker(self.files, out_path)
        self.worker.progress.connect(self.on_progress)
        self.worker.finished_all.connect(self.on_finished_all)
        self.worker.start()

    def on_progress(self, result: TaskResult):
        # Update progress bar & labels
        try:
            idx = self.files.index(result.src)
            pct = int((idx + 1) / max(1, len(self.files)) * 100)
            self.progress_bar.setValue(pct)
        except ValueError:
            # File not in list, shouldn't happen but handle gracefully
            pass

        # Show current processing file
        if idx < len(self.files) - 1:  # Not the last file
            self.lbl_processing.setText("正在处理：" + result.src.name)
        else:
            self.lbl_processing.setText("正在处理：完成")

        # Completed list (last up to 3)
        if not hasattr(self, "_completed_names"):
            self._completed_names = []
        if result.ok:
            self._completed_names.append(result.src.name)
            self._completed_names = self._completed_names[-3:]
            self.lbl_done.setText("已完成：" + ", ".join(self._completed_names))
        else:
            # Error list accumulate (last up to 3)
            if not hasattr(self, "_error_msgs"):
                self._error_msgs = []
            self._error_msgs.append(f"{result.src.name}: {result.message}")
            self._error_msgs = self._error_msgs[-3:]
            self.lbl_err.setText("错误：" + " | ".join(self._error_msgs))

    def on_finished_all(self):
        self.lbl_processing.setText("正在处理：全部完成")
        self.btn_start.setEnabled(True)
        self.btn_pick.setEnabled(True)
        self.btn_out.setEnabled(True)
        
        # Clean up worker
        if self.worker:
            self.worker.deleteLater()
            self.worker = None
            
        QtWidgets.QMessageBox.information(self, "完成", "处理完毕！")

    def clear_files(self):
        self.files = []
        self.update_selected_summary()


def main():
    # Windows scaling friendliness
    QtWidgets.QApplication.setAttribute(QtCore.Qt.AA_EnableHighDpiScaling, True)
    app = QtWidgets.QApplication(sys.argv)
    w = IndentorApp()
    w.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
